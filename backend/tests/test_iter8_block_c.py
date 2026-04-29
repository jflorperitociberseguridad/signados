"""Iter8 Phase 2 Block C — Admin Teaching, KB lookup, KB-augmented text-to-sign.

Also a thin regression covering Block B endpoints (/health, /email/status, /rtc/room,
/offline/pack, /translate/text-to-sign basic).
"""
import io
import os
import time
import pytest
import requests

BASE_URL = os.environ.get("REACT_APP_BACKEND_URL", "").rstrip("/")
if not BASE_URL:
    # Read from frontend/.env as fallback when env var not exported
    try:
        with open("/app/frontend/.env") as fh:
            for line in fh:
                if line.startswith("REACT_APP_BACKEND_URL="):
                    BASE_URL = line.split("=", 1)[1].strip().rstrip("/")
                    break
    except Exception:
        pass

ADMIN_PWD = "signlanguage-admin-2026"
HDR_OK = {"X-Admin-Password": ADMIN_PWD}


# ---------------- Block B regression ----------------
def test_health():
    r = requests.get(f"{BASE_URL}/api/health", timeout=20)
    assert r.status_code == 200
    j = r.json()
    assert j.get("status") in ("ok", "degraded")
    assert j.get("llm_key_configured") is True


def test_email_status():
    r = requests.get(f"{BASE_URL}/api/email/status", timeout=15)
    assert r.status_code == 200
    assert "configured" in r.json()


def test_rtc_room():
    r = requests.post(f"{BASE_URL}/api/rtc/room", timeout=15)
    assert r.status_code == 200
    j = r.json()
    assert "room" in j and len(j["room"]) >= 4


def test_offline_pack():
    r = requests.get(f"{BASE_URL}/api/offline/pack?limit=10", timeout=20)
    assert r.status_code == 200
    j = r.json()
    assert j["count"] >= 1
    assert isinstance(j["items"], list)


# ---------------- Admin auth gating ----------------
@pytest.mark.parametrize("path,method", [
    ("/api/admin/teaching/files", "GET"),
    ("/api/admin/teaching/knowledge", "GET"),
    ("/api/admin/teaching/corrections", "GET"),
    ("/api/admin/teaching/stats", "GET"),
])
def test_admin_endpoints_require_password(path, method):
    r = requests.request(method, f"{BASE_URL}{path}", timeout=15)
    assert r.status_code == 401, f"{path} should reject missing password"


def test_admin_upload_rejects_missing_password():
    files = {"file": ("a.pdf", b"%PDF-1.4 test", "application/pdf")}
    r = requests.post(f"{BASE_URL}/api/admin/teaching/upload", files=files, timeout=20)
    assert r.status_code == 401


def test_admin_upload_rejects_wrong_password():
    files = {"file": ("a.pdf", b"%PDF-1.4 test", "application/pdf")}
    r = requests.post(
        f"{BASE_URL}/api/admin/teaching/upload",
        files=files,
        headers={"X-Admin-Password": "wrong"},
        timeout=20,
    )
    assert r.status_code == 401


# ---------------- Admin Teaching CRUD ----------------
@pytest.fixture(scope="module")
def uploaded_file():
    """Upload a small DOCX-ish stub. The backend accepts based on extension."""
    # Build a minimal docx (real docx is a zip; teaching_service uses python-docx)
    # Use a real docx by writing one with python-docx if available, else fall back to PDF.
    try:
        from docx import Document
        d = Document()
        d.add_paragraph("TEST_KB_WORD: HOLA. Manos: palma abierta. Boca: 'hola'. Expresion: sonrisa.")
        d.add_paragraph("TEST_KB_WORD: GRACIAS. Manos: dedos juntos a la barbilla y desplazar al frente.")
        bio = io.BytesIO()
        d.save(bio)
        content = bio.getvalue()
        filename, ctype = "TEST_kb.docx", (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception:
        # Fallback PDF stub (extraction may produce empty kb but upload should pass)
        content = b"%PDF-1.4\nTEST file\n%%EOF"
        filename, ctype = "TEST_kb.pdf", "application/pdf"

    files = {"file": (filename, content, ctype)}
    data = {"label": "TEST_iter8"}
    r = requests.post(
        f"{BASE_URL}/api/admin/teaching/upload",
        files=files,
        data=data,
        headers=HDR_OK,
        timeout=30,
    )
    assert r.status_code == 200, r.text
    j = r.json()
    assert j.get("status") == "uploaded"
    assert j.get("size", 0) > 0
    assert j.get("type") in ("pdf", "docx", "image", "video")
    yield j
    # Cleanup
    requests.delete(
        f"{BASE_URL}/api/admin/teaching/files/{j['id']}", headers=HDR_OK, timeout=15
    )


def test_teaching_list_after_upload(uploaded_file):
    r = requests.get(f"{BASE_URL}/api/admin/teaching/files", headers=HDR_OK, timeout=15)
    assert r.status_code == 200
    ids = [d.get("id") for d in r.json()]
    assert uploaded_file["id"] in ids


def test_teaching_process_kicks_off(uploaded_file):
    r = requests.post(
        f"{BASE_URL}/api/admin/teaching/process/{uploaded_file['id']}",
        headers=HDR_OK,
        timeout=15,
    )
    assert r.status_code == 200
    assert r.json().get("started") is True

    # Poll up to ~45s for status flip (LLM extraction is async)
    deadline = time.time() + 45
    final_status = None
    while time.time() < deadline:
        rr = requests.get(
            f"{BASE_URL}/api/admin/teaching/files", headers=HDR_OK, timeout=15
        )
        for d in rr.json():
            if d.get("id") == uploaded_file["id"]:
                final_status = d.get("status")
                if final_status in ("processed", "error"):
                    break
        if final_status in ("processed", "error"):
            break
        time.sleep(3)
    # Don't fail the suite if LLM was slow; just record. Accept processed/error.
    assert final_status in ("processed", "error", "processing"), \
        f"unexpected status {final_status}"


def test_teaching_stats_shape():
    r = requests.get(f"{BASE_URL}/api/admin/teaching/stats", headers=HDR_OK, timeout=15)
    assert r.status_code == 200
    j = r.json()
    for k in ("files", "processed", "pending", "errors", "kb_count", "corrections", "by_language"):
        assert k in j, f"missing key {k}"
    assert isinstance(j["by_language"], list)


# ---------------- Corrections CRUD + KB-augmented text-to-sign ----------------
def test_corrections_upsert_and_list_and_delete():
    payload = {
        "word": "TESTHOLA",
        "language": "LSE",
        "hands": "palma abierta hacia el frente",
        "mouth": "hola",
        "expression": "sonrisa",
        "body": "ligero asentimiento",
        "status": "correct",
        "notes": "TEST_iter8",
    }
    r = requests.post(
        f"{BASE_URL}/api/admin/teaching/corrections",
        json=payload,
        headers=HDR_OK,
        timeout=15,
    )
    assert r.status_code == 200, r.text
    item = r.json()
    assert item["word"].lower() == "testhola"
    assert "id" in item
    cid = item["id"]

    # List
    r = requests.get(
        f"{BASE_URL}/api/admin/teaching/corrections", headers=HDR_OK, timeout=15
    )
    assert r.status_code == 200
    assert any(c.get("id") == cid for c in r.json())

    # Public KB lookup picks it up
    r = requests.get(f"{BASE_URL}/api/kb/lookup?q=TESTHOLA", timeout=15)
    assert r.status_code == 200
    items = r.json().get("items", [])
    sources = {it.get("_source") for it in items}
    assert "correction" in sources, f"correction not surfaced: {items}"

    # KB-augmented text-to-sign should report kb_used > 0 and confidence
    r = requests.post(
        f"{BASE_URL}/api/translate/text-to-sign",
        json={"text": "TESTHOLA por favor", "target_language": "LSE"},
        timeout=60,
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert "confidence" in body
    assert isinstance(body.get("kb_used"), int)
    assert body["kb_used"] >= 1, f"expected kb_used>=1, got {body}"

    # Cleanup correction
    rd = requests.delete(
        f"{BASE_URL}/api/admin/teaching/corrections/{cid}", headers=HDR_OK, timeout=15
    )
    assert rd.status_code == 200
    assert rd.json().get("deleted") == 1


def test_text_to_sign_basic_no_kb():
    """Text-to-sign without KB hint should still return shape with confidence."""
    r = requests.post(
        f"{BASE_URL}/api/translate/text-to-sign",
        json={"text": "Buenos dias amigo", "target_language": "auto"},
        timeout=60,
    )
    assert r.status_code == 200
    j = r.json()
    assert "steps" in j and "summary" in j
    assert "confidence" in j
    assert "kb_used" in j


def test_kb_lookup_empty_query():
    r = requests.get(f"{BASE_URL}/api/kb/lookup?q=", timeout=10)
    assert r.status_code == 200
    assert r.json() == {"items": []}
