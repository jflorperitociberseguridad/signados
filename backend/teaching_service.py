"""
Teaching service — admin-only knowledge ingestion pipeline.

Handles uploads of PDF/DOCX/IMAGE/VIDEO files, extracts text/frames,
and uses GPT-4o (vision) / GPT-4o-mini (text) to mine sign-language
knowledge into a structured KB stored in MongoDB.

Collections:
- teaching_files       { id, filename, type, size, status, uploaded_at, processed_at, kb_count, error }
- knowledge_base       { id, source_file_id, word, language, hands, mouth, expression, body, examples, confidence, source_type }
- corrections          { id, word, language, description, hands, mouth, expression, status, notes, admin, created_at, updated_at }
"""
from __future__ import annotations

import asyncio
import base64
import io
import json
import logging
import os
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional, Tuple

import cv2  # type: ignore
from pypdf import PdfReader  # type: ignore

logger = logging.getLogger("signlanguage.teaching")


SUPPORTED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "docx",
    "image/jpeg": "image",
    "image/jpg": "image",
    "image/png": "image",
    "image/webp": "image",
    "video/mp4": "video",
    "video/webm": "video",
    "video/quicktime": "video",
}


def teaching_dir() -> Path:
    p = Path(os.environ.get("TEACHING_DIR") or "/app/backend/data/teaching")
    p.mkdir(parents=True, exist_ok=True)
    return p


def detect_type(content_type: str, filename: str) -> Optional[str]:
    mime = (content_type or "").lower()
    if mime in SUPPORTED_TYPES:
        return SUPPORTED_TYPES[mime]
    fn = (filename or "").lower()
    for ext, t in {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
        ".webp": "image",
        ".mp4": "video",
        ".webm": "video",
        ".mov": "video",
    }.items():
        if fn.endswith(ext):
            return t
    return None


def safe_name(name: str) -> str:
    keep = "".join(c if c.isalnum() or c in ("-", "_", ".") else "_" for c in (name or "file"))
    return keep[:120] or "file"


# ---------------------------------------------------------------------------
# Extractors (sync, must run via asyncio.to_thread)
# ---------------------------------------------------------------------------
def _extract_pdf_text(path: Path) -> str:
    try:
        reader = PdfReader(str(path))
        chunks = []
        for page in reader.pages[:80]:  # safety cap
            try:
                chunks.append(page.extract_text() or "")
            except Exception:
                pass
        return "\n\n".join(chunks).strip()
    except Exception as exc:
        logger.warning("pdf extract failed: %s", exc)
        return ""


def _extract_docx_text(path: Path) -> str:
    try:
        # Avoid the heavy `python-docx` dependency footprint at import-time
        from docx import Document  # type: ignore

        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        parts.append(cell.text.strip())
        return "\n".join(parts).strip()
    except Exception as exc:
        logger.warning("docx extract failed: %s", exc)
        return ""


def _image_b64(path: Path, max_dim: int = 720) -> Optional[str]:
    try:
        img = cv2.imread(str(path))
        if img is None:
            return None
        h, w = img.shape[:2]
        scale = min(1.0, max_dim / max(h, w))
        if scale < 1.0:
            img = cv2.resize(img, (int(w * scale), int(h * scale)))
        ok, buf = cv2.imencode(".jpg", img, [cv2.IMWRITE_JPEG_QUALITY, 78])
        if not ok:
            return None
        return base64.b64encode(buf.tobytes()).decode("ascii")
    except Exception:
        return None


def _video_frames_b64(path: Path, n: int = 6, max_dim: int = 720) -> List[str]:
    try:
        cap = cv2.VideoCapture(str(path))
        if not cap.isOpened():
            return []
        total = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
        idxs = [int(total * i / n) for i in range(n)] if total else list(range(n))
        out = []
        for idx in idxs:
            if total:
                cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
            ok, frame = cap.read()
            if not ok:
                continue
            h, w = frame.shape[:2]
            scale = min(1.0, max_dim / max(h, w))
            if scale < 1.0:
                frame = cv2.resize(frame, (int(w * scale), int(h * scale)))
            ok2, buf = cv2.imencode(".jpg", frame, [cv2.IMWRITE_JPEG_QUALITY, 78])
            if ok2:
                out.append(base64.b64encode(buf.tobytes()).decode("ascii"))
        cap.release()
        return out
    except Exception:
        return []


# ---------------------------------------------------------------------------
# IA mining (async)
# ---------------------------------------------------------------------------
KB_SYSTEM_PROMPT = """Eres un asistente experto que extrae conocimiento estructurado de lengua de signos a partir de manuales, vídeos e imágenes pedagógicas.

Tu tarea: identifica todos los SIGNOS / SEÑAS / VOCABULARIO descritos en el material y devuélvelos en JSON.

Para cada signo extrae:
- word: palabra o concepto en español
- language: 'LSE', 'ASL', 'BSL', 'LSM', 'LIBRAS', 'Otro' o 'Desconocido'
- hands: descripción concisa de la configuración y movimiento de manos/dedos
- mouth: componente oral (labios/boca)
- expression: expresión facial (cejas, ojos, mejillas)
- body: postura del torso, hombros y movimiento corporal si aplica
- examples: 1-3 ejemplos prácticos o frases en las que se usa
- confidence: 'alta'|'media'|'baja' según lo claro que esté en la fuente

Devuelve EXCLUSIVAMENTE JSON válido (sin markdown, sin explicaciones extra) con esta forma:
{"items": [ { ...campos arriba... }, ... ]}

Si el material no contiene información identificable de signos, devuelve {"items": []}.
"""


def _parse_kb_json(raw: str) -> List[dict]:
    text = (raw or "").strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.lower().startswith("json"):
            text = text[4:].strip()
    try:
        obj = json.loads(text)
    except Exception:
        s, e = text.find("{"), text.rfind("}")
        if s != -1 and e > s:
            try:
                obj = json.loads(text[s : e + 1])
            except Exception:
                return []
        else:
            return []
    items = obj.get("items") if isinstance(obj, dict) else None
    if not isinstance(items, list):
        return []
    cleaned = []
    for it in items[:200]:
        if not isinstance(it, dict):
            continue
        word = (it.get("word") or "").strip()
        if not word:
            continue
        cleaned.append(
            {
                "word": word[:120],
                "language": (it.get("language") or "Desconocido")[:24],
                "hands": (it.get("hands") or "")[:600],
                "mouth": (it.get("mouth") or "")[:400],
                "expression": (it.get("expression") or "")[:400],
                "body": (it.get("body") or "")[:400],
                "examples": [str(x)[:200] for x in (it.get("examples") or [])[:5]],
                "confidence": (it.get("confidence") or "media").strip().lower()[:8],
            }
        )
    return cleaned


async def mine_with_llm(
    *,
    llm_chat_factory,  # callable(system, model) -> LlmChat
    user_message_factory,  # callable(text=..., file_contents=...) -> UserMessage
    image_content_cls,
    text_model: str,
    vision_model: str,
    text_chunks: List[str],
    images_b64: List[str],
) -> List[dict]:
    """Call the LLM with the available material and return cleaned KB entries.

    `text_chunks` and `images_b64` may be empty depending on file type.
    """
    if not text_chunks and not images_b64:
        return []

    items: List[dict] = []

    # ----- Text path (PDF / DOCX) -----
    if text_chunks:
        for chunk in text_chunks[:6]:  # safety cap
            chat = llm_chat_factory(KB_SYSTEM_PROMPT, text_model)
            prompt = (
                "Extrae signos del siguiente material y devuélvelos en JSON estricto:\n\n"
                + chunk[:14000]
            )
            try:
                raw = await chat.send_message(user_message_factory(text=prompt))
                items.extend(_parse_kb_json(raw))
            except Exception as exc:
                logger.warning("LLM text extract failed: %s", exc)

    # ----- Vision path (image/video frames) -----
    if images_b64:
        # 6 frames per call max
        for i in range(0, len(images_b64), 6):
            batch = images_b64[i : i + 6]
            chat = llm_chat_factory(KB_SYSTEM_PROMPT, vision_model)
            try:
                raw = await chat.send_message(
                    user_message_factory(
                        text=(
                            "Estos fotogramas/imagen son material pedagógico de "
                            "lengua de signos. Identifica los signos visibles y "
                            "devuelve JSON estricto."
                        ),
                        file_contents=[image_content_cls(image_base64=b) for b in batch],
                    )
                )
                items.extend(_parse_kb_json(raw))
            except Exception as exc:
                logger.warning("LLM vision extract failed: %s", exc)

    # Deduplicate by (word, language)
    seen = set()
    deduped = []
    for it in items:
        key = (it["word"].lower(), it["language"].lower())
        if key in seen:
            continue
        seen.add(key)
        deduped.append(it)
    return deduped


# ---------------------------------------------------------------------------
# Process pipeline (called by the API)
# ---------------------------------------------------------------------------
async def extract_material(file_path: Path, file_type: str) -> Tuple[List[str], List[str]]:
    """Return (text_chunks, images_b64) for the file."""
    text_chunks: List[str] = []
    images_b64: List[str] = []

    if file_type == "pdf":
        text = await asyncio.to_thread(_extract_pdf_text, file_path)
        if text:
            # Split into ~10k char chunks to fit in the context window
            CHUNK = 10000
            for i in range(0, len(text), CHUNK):
                text_chunks.append(text[i : i + CHUNK])
    elif file_type == "docx":
        text = await asyncio.to_thread(_extract_docx_text, file_path)
        if text:
            CHUNK = 10000
            for i in range(0, len(text), CHUNK):
                text_chunks.append(text[i : i + CHUNK])
    elif file_type == "image":
        b64 = await asyncio.to_thread(_image_b64, file_path)
        if b64:
            images_b64.append(b64)
    elif file_type == "video":
        frames = await asyncio.to_thread(_video_frames_b64, file_path, 8)
        images_b64.extend(frames)

    return text_chunks, images_b64
